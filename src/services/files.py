import os
import io
from pathlib import Path
from datetime import datetime
from src.config import get_settings

# Data directory for storing uploaded files
DATA_DIR = Path(__file__).resolve().parent.parent.parent / "data"

# Supported file extensions
SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}


class FileService:
    """Service for managing uploaded files."""
    
    def __init__(self):
        # Ensure data directory exists
        DATA_DIR.mkdir(exist_ok=True)
    
    def list_files(self) -> list[dict]:
        """List all uploaded files with metadata."""
        files = []
        for filepath in DATA_DIR.iterdir():
            if filepath.is_file() and filepath.name != ".gitkeep":
                stat = filepath.stat()
                files.append({
                    "filename": filepath.name,
                    "size_bytes": stat.st_size,
                    "uploaded_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "file_type": filepath.suffix.lower()
                })
        # Sort by upload time (newest first)
        files.sort(key=lambda x: x["uploaded_at"], reverse=True)
        return files
    
    def save_file(self, filename: str, content: bytes) -> dict:
        """Save uploaded file to data directory."""
        # Sanitize filename and make unique if exists
        safe_name = self._sanitize_filename(filename)
        filepath = DATA_DIR / safe_name
        
        # If file exists, add timestamp to make unique
        if filepath.exists():
            name, ext = os.path.splitext(safe_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = f"{name}_{timestamp}{ext}"
            filepath = DATA_DIR / safe_name
        
        # Write file
        filepath.write_bytes(content)
        
        return {
            "filename": safe_name,
            "size_bytes": len(content),
            "path": str(filepath),
            "file_type": filepath.suffix.lower()
        }
    
    def get_file(self, filename: str) -> tuple[bytes, str] | None:
        """Get file content and path. Returns None if not found."""
        filepath = DATA_DIR / filename
        if filepath.exists() and filepath.is_file():
            return filepath.read_bytes(), str(filepath)
        return None
    
    def get_file_text(self, filename: str) -> str | None:
        """
        Get text content from a file, supporting multiple formats.
        Returns None if file not found.
        """
        filepath = DATA_DIR / filename
        if not filepath.exists() or not filepath.is_file():
            return None
        
        ext = filepath.suffix.lower()
        
        if ext == ".txt":
            return filepath.read_text(encoding="utf-8")
        elif ext == ".docx":
            return self._extract_docx_text(filepath)
        elif ext == ".pdf":
            return self._extract_pdf_text(filepath)
        else:
            # Try reading as text for unknown formats
            try:
                return filepath.read_text(encoding="utf-8")
            except:
                return None
    
    def _extract_docx_text(self, filepath: Path) -> str:
        """Extract text content from a Word document with smart chunking."""
        try:
            from docx import Document
            doc = Document(filepath)
            
            # Collect all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Smart chunking: combine small paragraphs into meaningful chunks
            chunks = []
            current_chunk = []
            current_length = 0
            min_chunk_size = 100  # Minimum characters per chunk
            
            for para in paragraphs:
                current_chunk.append(para)
                current_length += len(para)
                
                # If we have enough content or hit a section break (numbered item, heading)
                is_section_start = (
                    para[0].isdigit() if para else False
                ) or para.isupper()
                
                if current_length >= min_chunk_size or is_section_start:
                    if current_chunk:
                        chunks.append(" ".join(current_chunk))
                        current_chunk = []
                        current_length = 0
            
            # Don't forget remaining content
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            
            # Join chunks with double newlines for later splitting
            return "\n\n".join(chunks)
        except Exception as e:
            print(f"Error extracting text from {filepath}: {e}")
            return ""
    
    def _extract_pdf_text(self, filepath: Path) -> str:
        """Extract text content from a PDF document."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(filepath)
            all_text = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text.strip())
            
            # Join pages and apply smart chunking
            full_text = "\n\n".join(all_text)
            
            # Split by paragraphs (double newlines) and apply chunking
            paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
            
            # Smart chunking
            chunks = []
            current_chunk = []
            current_length = 0
            min_chunk_size = 150
            
            for para in paragraphs:
                current_chunk.append(para)
                current_length += len(para)
                
                # Check for section breaks
                is_section_start = (
                    para[0].isdigit() if para else False
                ) or para.isupper() or para.startswith("•") or para.startswith("-")
                
                if current_length >= min_chunk_size or is_section_start:
                    if current_chunk:
                        chunks.append(" ".join(current_chunk))
                        current_chunk = []
                        current_length = 0
            
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            
            return "\n\n".join(chunks)
        except Exception as e:
            print(f"Error extracting text from PDF {filepath}: {e}")
            return ""
    
    def _extract_pdf_from_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            all_text = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text.strip())
            
            full_text = "\n\n".join(all_text)
            paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
            
            # Smart chunking
            chunks = []
            current_chunk = []
            current_length = 0
            min_chunk_size = 150
            
            for para in paragraphs:
                current_chunk.append(para)
                current_length += len(para)
                
                is_section_start = (
                    para[0].isdigit() if para else False
                ) or para.isupper()
                
                if current_length >= min_chunk_size or is_section_start:
                    if current_chunk:
                        chunks.append(" ".join(current_chunk))
                        current_chunk = []
                        current_length = 0
            
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            
            return "\n\n".join(chunks)
        except Exception as e:
            print(f"Error extracting text from PDF bytes: {e}")
            return ""
    
    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """Extract text from file bytes based on file extension."""
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".txt":
            return content.decode("utf-8")
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                
                # Collect all paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)
                
                # Smart chunking: combine small paragraphs
                chunks = []
                current_chunk = []
                current_length = 0
                min_chunk_size = 100
                
                for para in paragraphs:
                    current_chunk.append(para)
                    current_length += len(para)
                    
                    is_section_start = (
                        para[0].isdigit() if para else False
                    ) or para.isupper()
                    
                    if current_length >= min_chunk_size or is_section_start:
                        if current_chunk:
                            chunks.append(" ".join(current_chunk))
                            current_chunk = []
                            current_length = 0
                
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                
                return "\n\n".join(chunks)
            except Exception as e:
                print(f"Error extracting text from docx: {e}")
                return ""
        elif ext == ".pdf":
            return self._extract_pdf_from_bytes(content)
        else:
            try:
                return content.decode("utf-8")
            except:
                return ""
    
    def delete_file(self, filename: str) -> bool:
        """Delete a file. Returns True if deleted, False if not found."""
        filepath = DATA_DIR / filename
        if filepath.exists() and filepath.is_file():
            filepath.unlink()
            return True
        return False
    
    def file_exists(self, filename: str) -> bool:
        """Check if a file exists."""
        filepath = DATA_DIR / filename
        return filepath.exists() and filepath.is_file()
    
    def get_file_path(self, filename: str) -> Path | None:
        """Get the full path to a file."""
        filepath = DATA_DIR / filename
        if filepath.exists():
            return filepath
        return None
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if the file type is supported."""
        ext = os.path.splitext(filename)[1].lower()
        return ext in SUPPORTED_EXTENSIONS
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent directory traversal attacks."""
        # Remove any path components, keep only the filename
        return os.path.basename(filename).replace(" ", "_")


